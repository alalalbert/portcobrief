import csv
import os
from docx import Document
from scrape_page_bs import scrape_and_summarize
import concurrent.futures
import json
import signal
import sys
import time

def read_urls_from_file(file_path):
    with open(file_path, 'r') as file:
        return [line.strip() for line in file if line.strip()]

def append_short_summary_to_csv(company_name, url, short_summary, output_file):
    file_exists = os.path.isfile(output_file)
    with open(output_file, 'a', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        if not file_exists:
            writer.writerow(['Company Name', 'URL', 'Short Summary'])
        writer.writerow([company_name, url, short_summary])

def append_long_summary_to_docx(url, long_summary, output_file):
    if os.path.exists(output_file):
        doc = Document(output_file)
    else:
        doc = Document()
    
    doc.add_heading(url, level=1)
    doc.add_paragraph(long_summary)
    doc.add_paragraph()  # Add a blank line between summaries
    doc.save(output_file)

def scrape_url(url):
    print(f"Scraping: {url}")
    try:
        return url, scrape_and_summarize(url)
    except Exception as e:
        print(f"Error scraping {url}: {str(e)}")
        return url, None

def load_progress(progress_file):
    if os.path.exists(progress_file):
        with open(progress_file, 'r') as f:
            return set(json.load(f))
    return set()

def save_progress(progress_file, processed_urls):
    with open(progress_file, 'w') as f:
        json.dump(list(processed_urls), f)

def signal_handler(signum, frame):
    print("\nInterrupt received, stopping...")
    sys.exit(0)

def main():
    if len(sys.argv) != 2:
        print("Usage: python main.py <vc_portfolio_url>")
        sys.exit(1)
        
    vc_portfolio_url = sys.argv[1]
    
    # First, scrape the portfolio page to get company URLs
    from portcopage_scrape import scrape_portfolio_page
    print(f"Scraping portfolio page: {vc_portfolio_url}")
    scrape_portfolio_page(vc_portfolio_url)
    
    # Set up the signal handler
    signal.signal(signal.SIGINT, signal_handler)
    
    url_file = "company_urls.txt"
    progress_file = "progress.json"
    short_summaries_file = "short_summaries.csv"
    long_summaries_file = "long_summaries.docx"
    
    urls = read_urls_from_file(url_file)
    processed_urls = load_progress(progress_file)
    urls_to_process = [url for url in urls if url not in processed_urls]
    
    print(f"Processing {len(urls_to_process)} URLs")

    try:
        for url in urls_to_process:
            try:
                company_name, long_summary, short_summary, _ = scrape_and_summarize(url)
                
                if long_summary != "No data could be scraped." and long_summary != "Error scraping data:":
                    append_short_summary_to_csv(company_name, url, short_summary, short_summaries_file)
                    append_long_summary_to_docx(url, long_summary, long_summaries_file)
                    processed_urls.add(url)
                    save_progress(progress_file, processed_urls)
                    print(f"Processed and saved summaries for {url}")
                else:
                    print(f"Skipping {url} due to scraping error")
            except Exception as e:
                print(f"Error processing {url}: {str(e)}")
            
            # Add a small delay between requests to avoid overwhelming the server
            time.sleep(1)
    
    except KeyboardInterrupt:
        print("\nInterrupt received, stopping...")
    
    print("Processing complete. Results saved in short_summaries.csv and long_summaries.docx")

if __name__ == "__main__":
    main()
